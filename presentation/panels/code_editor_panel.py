# Code Editor Panel - Multi-tab Code Editor with Syntax Highlighting
"""
代码编辑器面板 - 多标签页代码编辑器

职责：
- 显示和编辑文件内容，支持多种格式
- 多标签页管理，支持拖拽排序
- 语法高亮（SPICE、JSON）
- 文档预览（Markdown、Word、PDF、图片）

位置：中栏（60%宽度）

设计原则：
- 延迟获取 ServiceLocator 中的服务
- 实现 retranslate_ui() 方法支持语言切换
- 订阅项目打开/关闭事件响应项目切换
- 文件操作通过 FileManager 进行
"""

import os
import re
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QTabWidget, QTabBar,
    QPlainTextEdit, QTextEdit, QLabel, QScrollArea, QMenu,
    QApplication, QSplitter, QFrame, QMessageBox, QPushButton,
    QToolButton
)
from PyQt6.QtCore import (
    Qt, pyqtSignal, QRect, QSize, QPoint, QTimer
)
from PyQt6.QtGui import (
    QFont, QColor, QPainter, QTextFormat, QSyntaxHighlighter,
    QTextCharFormat, QTextDocument, QAction, QPixmap, QImage,
    QPalette, QKeySequence, QShortcut, QTextCursor
)


# ============================================================
# 文件类型常量
# ============================================================

# 可编辑的文本文件扩展名
EDITABLE_EXTENSIONS = {'.cir', '.sp', '.spice', '.json', '.txt'}

# 图片文件扩展名
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp'}

# 文档文件扩展名（只读预览）
DOCUMENT_EXTENSIONS = {'.md', '.markdown', '.docx', '.pdf'}


# ============================================================
# 语法高亮器 - SPICE
# ============================================================

class SpiceHighlighter(QSyntaxHighlighter):
    """
    SPICE 文件语法高亮器
    
    高亮规则：
    - 注释（* 开头）：绿色
    - 指令（. 开头）：蓝色
    - 元件名（R、C、L、Q、M、D 等）：青色
    - 数值和单位：浅绿
    - 节点名：黄色
    """
    
    def __init__(self, document: QTextDocument):
        super().__init__(document)
        self._rules: List[Tuple[re.Pattern, QTextCharFormat]] = []
        self._setup_rules()
    
    def _setup_rules(self):
        """设置高亮规则"""
        # 注释格式（* 开头的行）
        comment_format = QTextCharFormat()
        comment_format.setForeground(QColor("#6a9955"))
        self._rules.append((
            re.compile(r'^\*.*$', re.MULTILINE),
            comment_format
        ))
        
        # 行内注释（; 后面的内容）
        inline_comment_format = QTextCharFormat()
        inline_comment_format.setForeground(QColor("#6a9955"))
        self._rules.append((
            re.compile(r';.*$', re.MULTILINE),
            inline_comment_format
        ))
        
        # 指令格式（. 开头，如 .tran, .ac, .dc, .param, .subckt, .ends, .include, .lib）
        directive_format = QTextCharFormat()
        directive_format.setForeground(QColor("#569cd6"))
        directive_format.setFontWeight(QFont.Weight.Bold)
        self._rules.append((
            re.compile(r'^\s*\.[a-zA-Z]+', re.MULTILINE | re.IGNORECASE),
            directive_format
        ))
        
        # 元件名格式（R、C、L、Q、M、D、V、I、E、F、G、H、X 开头）
        component_format = QTextCharFormat()
        component_format.setForeground(QColor("#4ec9b0"))
        self._rules.append((
            re.compile(r'\b[RCLQMDVIFGHEX][a-zA-Z0-9_]*\b', re.IGNORECASE),
            component_format
        ))
        
        # 数值和单位格式
        number_format = QTextCharFormat()
        number_format.setForeground(QColor("#b5cea8"))
        self._rules.append((
            re.compile(r'\b\d+\.?\d*[a-zA-Z]*\b'),
            number_format
        ))
        
        # 科学计数法
        self._rules.append((
            re.compile(r'\b\d+\.?\d*[eE][+-]?\d+\b'),
            number_format
        ))
        
        # 字符串格式（引号内）
        string_format = QTextCharFormat()
        string_format.setForeground(QColor("#ce9178"))
        self._rules.append((
            re.compile(r'"[^"]*"'),
            string_format
        ))
        self._rules.append((
            re.compile(r"'[^']*'"),
            string_format
        ))
    
    def highlightBlock(self, text: str):
        """高亮文本块"""
        for pattern, fmt in self._rules:
            for match in pattern.finditer(text):
                start = match.start()
                length = match.end() - start
                self.setFormat(start, length, fmt)


# ============================================================
# 语法高亮器 - JSON
# ============================================================

class JsonHighlighter(QSyntaxHighlighter):
    """
    JSON 文件语法高亮器
    
    高亮规则：
    - 键名：浅蓝
    - 字符串值：橙色
    - 数值：浅绿
    - 布尔值/null：蓝色
    """
    
    def __init__(self, document: QTextDocument):
        super().__init__(document)
        self._rules: List[Tuple[re.Pattern, QTextCharFormat]] = []
        self._setup_rules()
    
    def _setup_rules(self):
        """设置高亮规则"""
        # 键名格式（引号内，后跟冒号）
        key_format = QTextCharFormat()
        key_format.setForeground(QColor("#9cdcfe"))
        self._rules.append((
            re.compile(r'"[^"]*"\s*:'),
            key_format
        ))
        
        # 字符串值格式
        string_format = QTextCharFormat()
        string_format.setForeground(QColor("#ce9178"))
        self._rules.append((
            re.compile(r':\s*"[^"]*"'),
            string_format
        ))
        
        # 数值格式
        number_format = QTextCharFormat()
        number_format.setForeground(QColor("#b5cea8"))
        self._rules.append((
            re.compile(r':\s*-?\d+\.?\d*([eE][+-]?\d+)?'),
            number_format
        ))
        
        # 布尔值和 null 格式
        keyword_format = QTextCharFormat()
        keyword_format.setForeground(QColor("#569cd6"))
        self._rules.append((
            re.compile(r'\b(true|false|null)\b'),
            keyword_format
        ))
    
    def highlightBlock(self, text: str):
        """高亮文本块"""
        for pattern, fmt in self._rules:
            for match in pattern.finditer(text):
                start = match.start()
                length = match.end() - start
                self.setFormat(start, length, fmt)


# ============================================================
# 行号区域
# ============================================================

class LineNumberArea(QWidget):
    """
    行号区域组件
    
    显示代码编辑器的行号
    """
    
    def __init__(self, editor: 'CodeEditor'):
        super().__init__(editor)
        self._editor = editor
    
    def sizeHint(self) -> QSize:
        return QSize(self._editor.line_number_area_width(), 0)
    
    def paintEvent(self, event):
        self._editor.line_number_area_paint_event(event)


# ============================================================
# 代码编辑器组件
# ============================================================

class CodeEditor(QPlainTextEdit):
    """
    代码编辑器组件
    
    功能：
    - 行号显示
    - 当前行高亮
    - 语法高亮
    - 基本快捷键支持
    
    信号：
    - modification_changed(bool): 修改状态变化时发出
    """
    
    # 修改状态变化信号
    modification_changed = pyqtSignal(bool)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        
        # 行号区域
        self._line_number_area = LineNumberArea(self)
        
        # 语法高亮器
        self._highlighter: Optional[QSyntaxHighlighter] = None
        
        # 文件路径
        self._file_path: Optional[str] = None
        
        # 是否已修改
        self._is_modified = False
        
        # 设置字体
        self._setup_font()
        
        # 设置样式
        self._setup_style()
        
        # 连接信号
        self.blockCountChanged.connect(self._update_line_number_area_width)
        self.updateRequest.connect(self._update_line_number_area)
        self.cursorPositionChanged.connect(self._highlight_current_line)
        self.textChanged.connect(self._on_text_changed)
        
        # 初始化
        self._update_line_number_area_width(0)
        self._highlight_current_line()
    
    def _setup_font(self):
        """设置编程字体"""
        font = QFont()
        # 尝试使用常见的编程字体
        for font_name in ["JetBrains Mono", "Consolas", "Fira Code", "Monaco", "Courier New"]:
            font.setFamily(font_name)
            if font.exactMatch():
                break
        font.setPointSize(11)
        font.setFixedPitch(True)
        self.setFont(font)
        
        # 设置 Tab 宽度为 4 个空格
        self.setTabStopDistance(self.fontMetrics().horizontalAdvance(' ') * 4)
    
    def _setup_style(self):
        """设置编辑器样式"""
        self.setStyleSheet("""
            QPlainTextEdit {
                background-color: #ffffff;
                color: #333333;
                border: none;
                selection-background-color: #e3f2fd;
                selection-color: #333333;
            }
        """)

    def line_number_area_width(self) -> int:
        """计算行号区域宽度"""
        digits = 1
        max_num = max(1, self.blockCount())
        while max_num >= 10:
            max_num //= 10
            digits += 1
        
        # 最少显示 3 位数字的宽度
        digits = max(3, digits)
        
        space = 10 + self.fontMetrics().horizontalAdvance('9') * digits
        return space
    
    def _update_line_number_area_width(self, _):
        """更新行号区域宽度"""
        self.setViewportMargins(self.line_number_area_width(), 0, 0, 0)
    
    def _update_line_number_area(self, rect: QRect, dy: int):
        """更新行号区域"""
        if dy:
            self._line_number_area.scroll(0, dy)
        else:
            self._line_number_area.update(
                0, rect.y(),
                self._line_number_area.width(), rect.height()
            )
        
        if rect.contains(self.viewport().rect()):
            self._update_line_number_area_width(0)
    
    def resizeEvent(self, event):
        """窗口大小变化"""
        super().resizeEvent(event)
        cr = self.contentsRect()
        self._line_number_area.setGeometry(
            QRect(cr.left(), cr.top(),
                  self.line_number_area_width(), cr.height())
        )
    
    def line_number_area_paint_event(self, event):
        """绘制行号区域"""
        painter = QPainter(self._line_number_area)
        painter.fillRect(event.rect(), QColor("#f8f9fa"))
        
        # 绘制右侧分隔线
        painter.setPen(QColor("#e0e0e0"))
        painter.drawLine(
            self._line_number_area.width() - 1, event.rect().top(),
            self._line_number_area.width() - 1, event.rect().bottom()
        )
        
        # 绘制行号
        block = self.firstVisibleBlock()
        block_number = block.blockNumber()
        top = int(self.blockBoundingGeometry(block).translated(
            self.contentOffset()).top())
        bottom = top + int(self.blockBoundingRect(block).height())
        
        painter.setPen(QColor("#999999"))
        
        while block.isValid() and top <= event.rect().bottom():
            if block.isVisible() and bottom >= event.rect().top():
                number = str(block_number + 1)
                painter.drawText(
                    0, top,
                    self._line_number_area.width() - 5,
                    self.fontMetrics().height(),
                    Qt.AlignmentFlag.AlignRight,
                    number
                )
            
            block = block.next()
            top = bottom
            bottom = top + int(self.blockBoundingRect(block).height())
            block_number += 1
    
    def _highlight_current_line(self):
        """高亮当前行"""
        extra_selections = []
        
        if not self.isReadOnly():
            selection = QTextEdit.ExtraSelection()
            line_color = QColor("#f0f7ff")
            selection.format.setBackground(line_color)
            selection.format.setProperty(
                QTextFormat.Property.FullWidthSelection, True
            )
            selection.cursor = self.textCursor()
            selection.cursor.clearSelection()
            extra_selections.append(selection)
        
        self.setExtraSelections(extra_selections)
    
    def _on_text_changed(self):
        """文本变化处理"""
        if not self._is_modified:
            self._is_modified = True
            # 发出修改状态变化信号
            self.modification_changed.emit(True)
    
    def set_highlighter(self, file_ext: str):
        """根据文件扩展名设置语法高亮器"""
        # 移除旧的高亮器
        if self._highlighter:
            self._highlighter.setDocument(None)
            self._highlighter = None
        
        # 根据扩展名创建新的高亮器
        ext = file_ext.lower()
        if ext in {'.cir', '.sp', '.spice'}:
            self._highlighter = SpiceHighlighter(self.document())
        elif ext == '.json':
            self._highlighter = JsonHighlighter(self.document())
    
    def get_file_path(self) -> Optional[str]:
        """获取文件路径"""
        return self._file_path
    
    def set_file_path(self, path: str):
        """设置文件路径"""
        self._file_path = path
    
    def is_modified(self) -> bool:
        """是否已修改"""
        return self._is_modified
    
    def set_modified(self, modified: bool):
        """设置修改状态"""
        old_modified = self._is_modified
        self._is_modified = modified
        # 状态变化时发出信号
        if old_modified != modified:
            self.modification_changed.emit(modified)
    
    def get_cursor_position(self) -> Tuple[int, int]:
        """获取光标位置（行号，列号）"""
        cursor = self.textCursor()
        return cursor.blockNumber() + 1, cursor.columnNumber() + 1


# ============================================================
# 图片预览组件
# ============================================================

class ImageViewer(QScrollArea):
    """
    图片预览组件
    
    功能：
    - 居中显示图片
    - 支持缩放
    """
    
    def __init__(self, parent=None):
        super().__init__(parent)
        
        # 图片标签
        self._image_label = QLabel()
        self._image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._image_label.setStyleSheet("background-color: #f5f5f5;")
        
        # 原始图片
        self._original_pixmap: Optional[QPixmap] = None
        
        # 缩放比例
        self._scale_factor = 1.0
        
        # 设置滚动区域
        self.setWidget(self._image_label)
        self.setWidgetResizable(True)
        self.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.setStyleSheet("background-color: #f5f5f5; border: none;")
    
    def load_image(self, path: str) -> bool:
        """加载图片"""
        pixmap = QPixmap(path)
        if pixmap.isNull():
            self._image_label.setText("Failed to load image")
            return False
        
        self._original_pixmap = pixmap
        self._scale_factor = 1.0
        self._update_display()
        return True
    
    def _update_display(self):
        """更新显示"""
        if self._original_pixmap is None:
            return
        
        # 计算缩放后的尺寸
        scaled_pixmap = self._original_pixmap.scaled(
            int(self._original_pixmap.width() * self._scale_factor),
            int(self._original_pixmap.height() * self._scale_factor),
            Qt.AspectRatioMode.KeepAspectRatio,
            Qt.TransformationMode.SmoothTransformation
        )
        
        self._image_label.setPixmap(scaled_pixmap)
    
    def zoom_in(self):
        """放大"""
        self._scale_factor *= 1.25
        self._update_display()
    
    def zoom_out(self):
        """缩小"""
        self._scale_factor *= 0.8
        self._update_display()
    
    def fit_to_window(self):
        """适应窗口"""
        if self._original_pixmap is None:
            return
        
        # 计算适应窗口的缩放比例
        viewport_size = self.viewport().size()
        img_size = self._original_pixmap.size()
        
        scale_w = viewport_size.width() / img_size.width()
        scale_h = viewport_size.height() / img_size.height()
        
        self._scale_factor = min(scale_w, scale_h, 1.0)
        self._update_display()


# ============================================================
# 文档预览组件
# ============================================================

class DocumentViewer(QTextEdit):
    """
    文档预览组件
    
    功能：
    - Markdown 渲染预览
    - Word 文档文本提取预览
    - PDF 文本提取预览
    """
    
    def __init__(self, parent=None):
        super().__init__(parent)
        
        # 设置只读
        self.setReadOnly(True)
        
        # 设置样式
        self.setStyleSheet("""
            QTextEdit {
                background-color: #ffffff;
                color: #333333;
                border: none;
                padding: 20px;
                font-family: "Segoe UI", "Microsoft YaHei", sans-serif;
                font-size: 14px;
                line-height: 1.6;
            }
        """)
    
    def load_markdown(self, path: str) -> bool:
        """加载 Markdown 文件"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 尝试使用 markdown 库渲染
            try:
                import markdown
                html = markdown.markdown(
                    content,
                    extensions=['tables', 'fenced_code', 'codehilite']
                )
                self.setHtml(self._wrap_html(html))
            except ImportError:
                # 如果没有 markdown 库，显示原始文本
                self.setPlainText(content)
            
            return True
        except Exception as e:
            self.setPlainText(f"Failed to load Markdown: {e}")
            return False
    
    def load_word(self, path: str) -> bool:
        """加载 Word 文档"""
        try:
            from docx import Document
            doc = Document(path)
            
            # 提取段落文本
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(f"<p>{para.text}</p>")
            
            html = "\n".join(paragraphs)
            self.setHtml(self._wrap_html(html))
            return True
        except ImportError:
            self.setPlainText("python-docx library not installed.\nInstall with: pip install python-docx")
            return False
        except Exception as e:
            self.setPlainText(f"Failed to load Word document: {e}")
            return False
    
    def load_pdf(self, path: str) -> bool:
        """加载 PDF 文档"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(path)
            
            # 提取所有页面的文本
            text_parts = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"<h3>Page {page_num + 1}</h3>")
                    text_parts.append(f"<pre>{text}</pre>")
            
            doc.close()
            
            html = "\n".join(text_parts)
            self.setHtml(self._wrap_html(html))
            return True
        except ImportError:
            self.setPlainText("PyMuPDF library not installed.\nInstall with: pip install PyMuPDF")
            return False
        except Exception as e:
            self.setPlainText(f"Failed to load PDF: {e}")
            return False
    
    def _wrap_html(self, content: str) -> str:
        """包装 HTML 内容"""
        return f"""
        <html>
        <head>
        <style>
            body {{ font-family: "Segoe UI", "Microsoft YaHei", sans-serif; line-height: 1.6; }}
            h1, h2, h3 {{ color: #333; }}
            code {{ background-color: #f5f5f5; padding: 2px 5px; border-radius: 3px; }}
            pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 5px; overflow-x: auto; }}
            table {{ border-collapse: collapse; width: 100%; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #f5f5f5; }}
        </style>
        </head>
        <body>
        {content}
        </body>
        </html>
        """


# ============================================================
# 标签页组件
# ============================================================

class EditorTab:
    """
    编辑器标签页数据
    
    存储每个标签页的状态信息
    """
    
    def __init__(self, path: str, widget: QWidget, is_readonly: bool = False):
        self.path = path
        self.widget = widget
        self.is_readonly = is_readonly
        self.is_modified = False


# ============================================================
# 自定义标签栏
# ============================================================

class EditorTabBar(QTabBar):
    """
    自定义标签栏
    
    功能：
    - 显示文件名和修改标记
    - 关闭按钮始终可见
    - 右键菜单
    - 悬停显示完整路径
    """
    
    # 关闭标签页信号
    tab_close_requested = pyqtSignal(int)
    
    def __init__(self, parent=None):
        super().__init__(parent)
        
        # 设置可关闭
        self.setTabsClosable(True)
        self.setMovable(True)
        self.setExpanding(False)
        
        # 连接关闭信号
        self.tabCloseRequested.connect(self.tab_close_requested.emit)
        
        # 获取关闭图标路径
        close_icon_path = self._get_close_icon_path()
        
        # 设置样式 - 关闭按钮始终可见
        self.setStyleSheet(f"""
            QTabBar {{
                background-color: #f8f9fa;
                border-bottom: 1px solid #e0e0e0;
            }}
            QTabBar::tab {{
                background-color: #f8f9fa;
                border: none;
                border-right: 1px solid #e0e0e0;
                padding: 6px 24px 6px 12px;
                min-width: 80px;
                max-width: 200px;
            }}
            QTabBar::tab:selected {{
                background-color: #ffffff;
                border-bottom: 2px solid #4CAF50;
            }}
            QTabBar::tab:hover:!selected {{
                background-color: #f0f7ff;
            }}
            QTabBar::close-button {{
                image: url({close_icon_path});
                subcontrol-position: right;
                subcontrol-origin: padding;
                width: 16px;
                height: 16px;
                margin-right: 4px;
                border-radius: 2px;
                background-color: transparent;
            }}
            QTabBar::close-button:hover {{
                background-color: #e0e0e0;
            }}
            QTabBar::close-button:pressed {{
                background-color: #d0d0d0;
            }}
        """)
    
    def _get_close_icon_path(self) -> str:
        """获取关闭图标路径"""
        try:
            from resources.resource_loader import get_icon_path
            path = get_icon_path("panel", "close")
            if path:
                # 将反斜杠转换为正斜杠（QSS 需要）
                return path.replace("\\", "/")
        except Exception:
            pass
        return ""


# ============================================================
# 代码编辑器面板
# ============================================================

class CodeEditorPanel(QWidget):
    """
    代码编辑器面板
    
    功能：
    - 多标签页管理
    - 支持多种文件格式
    - 语法高亮
    - 文档预览
    
    信号：
    - file_saved(str): 文件保存时发出，携带文件路径
    - open_workspace_requested(): 请求打开工作区时发出
    """
    
    # 文件保存信号
    file_saved = pyqtSignal(str)
    # 请求打开工作区信号
    open_workspace_requested = pyqtSignal()
    # 撤销/重做状态变化信号
    undo_redo_state_changed = pyqtSignal(bool, bool)  # (can_undo, can_redo)
    # 可编辑文件状态变化信号（用于启用/禁用保存按钮）
    editable_file_state_changed = pyqtSignal(bool)  # has_editable_file
    
    def __init__(self, parent=None):
        super().__init__(parent)
        
        # 延迟获取的服务
        self._i18n_manager = None
        self._event_bus = None
        self._file_manager = None
        self._logger = None
        
        # 标签页数据
        self._tabs: Dict[str, EditorTab] = {}
        
        # UI 组件
        self._tab_widget: Optional[QTabWidget] = None
        self._status_bar: Optional[QWidget] = None
        self._line_col_label: Optional[QLabel] = None
        self._encoding_label: Optional[QLabel] = None
        self._file_type_label: Optional[QLabel] = None
        self._readonly_label: Optional[QLabel] = None
        
        # 空状态占位
        self._empty_widget: Optional[QWidget] = None
        
        # 只读模式
        self._is_readonly_mode = False
        
        # 初始化 UI
        self._setup_ui()
        self._setup_shortcuts()
        
        # 应用国际化文本
        self.retranslate_ui()
        
        # 订阅事件
        self._subscribe_events()

    # ============================================================
    # 延迟获取服务
    # ============================================================

    @property
    def i18n_manager(self):
        """延迟获取 I18nManager"""
        if self._i18n_manager is None:
            try:
                from shared.service_locator import ServiceLocator
                from shared.service_names import SVC_I18N_MANAGER
                self._i18n_manager = ServiceLocator.get_optional(SVC_I18N_MANAGER)
            except Exception:
                pass
        return self._i18n_manager

    @property
    def event_bus(self):
        """延迟获取 EventBus"""
        if self._event_bus is None:
            try:
                from shared.service_locator import ServiceLocator
                from shared.service_names import SVC_EVENT_BUS
                self._event_bus = ServiceLocator.get_optional(SVC_EVENT_BUS)
            except Exception:
                pass
        return self._event_bus

    @property
    def file_manager(self):
        """延迟获取 FileManager"""
        if self._file_manager is None:
            try:
                from shared.service_locator import ServiceLocator
                from shared.service_names import SVC_FILE_MANAGER
                self._file_manager = ServiceLocator.get_optional(SVC_FILE_MANAGER)
            except Exception:
                pass
        return self._file_manager

    @property
    def logger(self):
        """延迟获取 Logger"""
        if self._logger is None:
            try:
                from infrastructure.utils.logger import get_logger
                self._logger = get_logger("code_editor_panel")
            except Exception:
                pass
        return self._logger

    def _get_text(self, key: str, default: Optional[str] = None) -> str:
        """获取国际化文本"""
        if self.i18n_manager:
            return self.i18n_manager.get_text(key, default)
        return default if default else key

    # ============================================================
    # UI 初始化
    # ============================================================

    def _setup_ui(self):
        """设置 UI 布局"""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # 创建标签栏区域（左按钮 + 标签栏 + 右按钮），固定高度
        tab_bar_container = QWidget()
        tab_bar_container.setFixedHeight(32)  # 固定标签栏高度
        tab_bar_layout = QHBoxLayout(tab_bar_container)
        tab_bar_layout.setContentsMargins(0, 0, 0, 0)
        tab_bar_layout.setSpacing(0)
        
        # 左侧滚动按钮（直接创建，不使用容器类）
        self._scroll_left_btn = QToolButton(tab_bar_container)
        self._scroll_left_btn.setArrowType(Qt.ArrowType.LeftArrow)
        self._scroll_left_btn.setFixedSize(24, 28)
        self._scroll_left_btn.setAutoRepeat(True)
        self._scroll_left_btn.setAutoRepeatInterval(100)
        self._scroll_left_btn.clicked.connect(self._on_scroll_left)
        self._scroll_left_btn.setStyleSheet("""
            QToolButton {
                background-color: #f0f0f0;
                border: 1px solid #d0d0d0;
                border-radius: 3px;
            }
            QToolButton:hover { background-color: #e0e0e0; }
            QToolButton:pressed { background-color: #d0d0d0; }
            QToolButton:disabled { background-color: #f8f8f8; border-color: #e0e0e0; }
        """)
        tab_bar_layout.addWidget(self._scroll_left_btn)
        
        # 标签页组件
        self._tab_widget = QTabWidget()
        # 先设置自定义标签栏
        self._tab_bar = EditorTabBar()
        self._tab_widget.setTabBar(self._tab_bar)
        # 禁用内置滚动按钮（我们使用自定义按钮）
        self._tab_widget.setUsesScrollButtons(False)
        # 设置标签文本省略模式（不省略，保持完整显示）
        self._tab_widget.setElideMode(Qt.TextElideMode.ElideNone)
        # 使用文档模式简化外观
        self._tab_widget.setDocumentMode(True)
        self._tab_widget.tabCloseRequested.connect(self._on_tab_close_requested)
        self._tab_widget.currentChanged.connect(self._on_current_tab_changed)
        self._tab_widget.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self._tab_widget.customContextMenuRequested.connect(self._on_tab_context_menu)
        
        # 将标签栏添加到布局（从 QTabWidget 中提取出来）
        tab_bar_layout.addWidget(self._tab_bar, 1)
        
        # 右侧滚动按钮
        self._scroll_right_btn = QToolButton(tab_bar_container)
        self._scroll_right_btn.setArrowType(Qt.ArrowType.RightArrow)
        self._scroll_right_btn.setFixedSize(24, 28)
        self._scroll_right_btn.setAutoRepeat(True)
        self._scroll_right_btn.setAutoRepeatInterval(100)
        self._scroll_right_btn.clicked.connect(self._on_scroll_right)
        self._scroll_right_btn.setStyleSheet("""
            QToolButton {
                background-color: #f0f0f0;
                border: 1px solid #d0d0d0;
                border-radius: 3px;
            }
            QToolButton:hover { background-color: #e0e0e0; }
            QToolButton:pressed { background-color: #d0d0d0; }
            QToolButton:disabled { background-color: #f8f8f8; border-color: #e0e0e0; }
        """)
        tab_bar_layout.addWidget(self._scroll_right_btn)
        
        # 初始隐藏滚动按钮（无标签时）
        self._scroll_left_btn.hide()
        self._scroll_right_btn.hide()
        
        layout.addWidget(tab_bar_container)
        
        # 标签页内容区域（QTabWidget 不含标签栏，因为标签栏已提取）
        layout.addWidget(self._tab_widget, 1)
        
        # 空状态占位
        self._empty_widget = self._create_empty_widget()
        layout.addWidget(self._empty_widget)
        
        # 底部状态条
        self._status_bar = self._create_status_bar()
        layout.addWidget(self._status_bar)
        
        # 初始显示空状态
        self._update_empty_state()

    def _create_empty_widget(self) -> QWidget:
        """创建空状态占位组件"""
        widget = QWidget()
        widget.setStyleSheet("background-color: #f5f5f5;")
        
        layout = QVBoxLayout(widget)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # 提示文本
        label = QLabel()
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        label.setStyleSheet("color: #888; font-size: 14px;")
        label.setProperty("empty_hint", True)
        layout.addWidget(label)
        
        # 打开工作区按钮（无项目时显示）
        self._open_workspace_btn = QPushButton()
        self._open_workspace_btn.setProperty("open_workspace_btn", True)
        self._open_workspace_btn.setFixedSize(200, 50)
        self._open_workspace_btn.setStyleSheet(
            "QPushButton { background-color: #4CAF50; color: white; "
            "border: none; border-radius: 8px; font-size: 16px; font-weight: bold; }"
            "QPushButton:hover { background-color: #45a049; }"
        )
        self._open_workspace_btn.clicked.connect(self._on_open_workspace_clicked)
        layout.addWidget(self._open_workspace_btn, alignment=Qt.AlignmentFlag.AlignCenter)
        
        return widget

    def _create_status_bar(self) -> QWidget:
        """创建底部状态条"""
        status_bar = QWidget()
        status_bar.setFixedHeight(24)
        status_bar.setStyleSheet("""
            QWidget {
                background-color: #f8f9fa;
                border-top: 1px solid #e0e0e0;
            }
            QLabel {
                color: #666666;
                font-size: 11px;
                padding: 0 8px;
            }
        """)
        
        layout = QHBoxLayout(status_bar)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        
        # 只读标签
        self._readonly_label = QLabel()
        self._readonly_label.setStyleSheet(
            "background-color: #ffeb3b; color: #333; padding: 2px 8px; border-radius: 2px;"
        )
        self._readonly_label.hide()
        layout.addWidget(self._readonly_label)
        
        layout.addStretch()
        
        # 行列号
        self._line_col_label = QLabel("Ln 1, Col 1")
        layout.addWidget(self._line_col_label)
        
        # 分隔符
        sep1 = QFrame()
        sep1.setFrameShape(QFrame.Shape.VLine)
        sep1.setStyleSheet("color: #e0e0e0;")
        layout.addWidget(sep1)
        
        # 编码
        self._encoding_label = QLabel("UTF-8")
        layout.addWidget(self._encoding_label)
        
        # 分隔符
        sep2 = QFrame()
        sep2.setFrameShape(QFrame.Shape.VLine)
        sep2.setStyleSheet("color: #e0e0e0;")
        layout.addWidget(sep2)
        
        # 文件类型
        self._file_type_label = QLabel("Plain Text")
        layout.addWidget(self._file_type_label)
        
        return status_bar

    def _setup_shortcuts(self):
        """设置快捷键"""
        # 注意：Ctrl+S 和 Ctrl+Shift+S 由菜单栏统一管理，避免快捷键冲突
        
        # Ctrl+W 关闭当前标签
        close_shortcut = QShortcut(QKeySequence("Ctrl+W"), self)
        close_shortcut.activated.connect(self._close_current_tab)

    def _update_empty_state(self):
        """更新空状态显示"""
        has_tabs = self._tab_widget.count() > 0
        self._tab_widget.setVisible(has_tabs)
        self._empty_widget.setVisible(not has_tabs)
        self._status_bar.setVisible(has_tabs)
        
        # 根据是否有项目显示/隐藏打开工作区按钮
        if hasattr(self, '_open_workspace_btn') and self._open_workspace_btn:
            has_project = self._check_has_project()
            self._open_workspace_btn.setVisible(not has_project)
            
            # 更新提示文本
            for child in self._empty_widget.findChildren(QLabel):
                if child.property("empty_hint"):
                    if has_project:
                        child.setText(self._get_text("hint.select_file", "Select a file to view"))
                    else:
                        child.setText(self._get_text("hint.open_workspace", "Open a workspace to get started"))
    
    def _check_has_project(self) -> bool:
        """检查是否有已打开的项目"""
        try:
            from shared.service_locator import ServiceLocator
            from shared.service_names import SVC_APP_STATE
            app_state = ServiceLocator.get_optional(SVC_APP_STATE)
            if app_state:
                from shared.app_state import STATE_PROJECT_PATH
                project_path = app_state.get(STATE_PROJECT_PATH)
                return project_path is not None and project_path != ""
        except Exception:
            pass
        return False
    
    def _on_open_workspace_clicked(self):
        """打开工作区按钮点击处理"""
        self.open_workspace_requested.emit()

    # ============================================================
    # 核心功能
    # ============================================================

    def load_file(self, path: str) -> bool:
        """
        加载文件内容
        
        Args:
            path: 文件路径
            
        Returns:
            bool: 是否成功
        """
        if not path or not os.path.isfile(path):
            if self.logger:
                self.logger.warning(f"Invalid file path: {path}")
            return False
        
        # 检查是否已打开
        if path in self._tabs:
            # 切换到已打开的标签页
            tab = self._tabs[path]
            index = self._tab_widget.indexOf(tab.widget)
            if index >= 0:
                self._tab_widget.setCurrentIndex(index)
            return True
        
        # 根据文件类型创建对应的组件
        ext = os.path.splitext(path)[1].lower()
        
        try:
            if ext in EDITABLE_EXTENSIONS:
                widget = self._create_code_editor(path, ext)
            elif ext in IMAGE_EXTENSIONS:
                widget = self._create_image_viewer(path)
            elif ext in DOCUMENT_EXTENSIONS:
                widget = self._create_document_viewer(path, ext)
            else:
                # 默认作为文本文件处理
                widget = self._create_code_editor(path, ext)
            
            if widget is None:
                return False
            
            # 添加标签页
            is_readonly = ext in IMAGE_EXTENSIONS or ext in DOCUMENT_EXTENSIONS
            self._add_tab(path, widget, is_readonly)
            
            if self.logger:
                self.logger.info(f"File loaded: {path}")
            
            return True
            
        except Exception as e:
            if self.logger:
                self.logger.error(f"Failed to load file: {path}, error: {e}")
            return False

    def _create_code_editor(self, path: str, ext: str) -> Optional[CodeEditor]:
        """创建代码编辑器"""
        editor = CodeEditor()
        editor.set_file_path(path)
        
        # 读取文件内容
        try:
            if self.file_manager:
                content = self.file_manager.read_file(path)
            else:
                with open(path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            # 阻止信号，避免设置内容和高亮器时触发修改状态变化
            editor.blockSignals(True)
            editor.document().blockSignals(True)
            
            editor.setPlainText(content)
            
            # 设置语法高亮（可能触发文档重新处理）
            editor.set_highlighter(ext)
            
            # 恢复信号
            editor.document().blockSignals(False)
            editor.blockSignals(False)
            
            # 彻底重置修改状态（同时重置 QTextDocument 和 CodeEditor 的状态）
            editor.document().setModified(False)
            editor.set_modified(False)
            
            # 连接光标位置变化信号
            editor.cursorPositionChanged.connect(self._update_cursor_position)
            
            # 连接撤销/重做状态变化信号
            editor.document().undoAvailable.connect(self._on_undo_available_changed)
            editor.document().redoAvailable.connect(self._on_redo_available_changed)
            
            # 连接修改状态变化信号（用于更新标签页标题）
            editor.modification_changed.connect(
                lambda modified, p=path: self._on_editor_modification_changed(p, modified)
            )
            
            # 设置只读模式
            if self._is_readonly_mode:
                editor.setReadOnly(True)
            
            return editor
            
        except Exception as e:
            if self.logger:
                self.logger.error(f"Failed to read file: {path}, error: {e}")
            return None

    def _create_image_viewer(self, path: str) -> Optional[ImageViewer]:
        """创建图片预览器"""
        viewer = ImageViewer()
        if viewer.load_image(path):
            viewer.fit_to_window()
            return viewer
        return None

    def _create_document_viewer(self, path: str, ext: str) -> Optional[DocumentViewer]:
        """创建文档预览器"""
        viewer = DocumentViewer()
        
        if ext in {'.md', '.markdown'}:
            if viewer.load_markdown(path):
                return viewer
        elif ext == '.docx':
            if viewer.load_word(path):
                return viewer
        elif ext == '.pdf':
            if viewer.load_pdf(path):
                return viewer
        
        return viewer  # 即使加载失败也返回，显示错误信息

    def _add_tab(self, path: str, widget: QWidget, is_readonly: bool):
        """添加标签页"""
        # 获取文件名
        file_name = os.path.basename(path)
        
        # 添加到标签页
        index = self._tab_widget.addTab(widget, file_name)
        self._tab_widget.setCurrentIndex(index)
        
        # 设置工具提示为完整路径
        self._tab_widget.setTabToolTip(index, path)
        
        # 保存标签页数据
        tab = EditorTab(path, widget, is_readonly)
        self._tabs[path] = tab
        
        # 更新空状态
        self._update_empty_state()
        
        # 更新状态栏
        self._update_status_bar(path)
        
        # 确保标签页标题正确（无修改标记）
        self._update_tab_title(path)
        
        # 通知可编辑文件状态变化（用于启用保存按钮）
        self._emit_editable_file_state()

    def save_file(self) -> bool:
        """
        保存当前文件
        
        Returns:
            bool: 是否成功
        """
        current_widget = self._tab_widget.currentWidget()
        if not current_widget:
            return False
        
        # 查找对应的标签页数据
        tab = None
        for t in self._tabs.values():
            if t.widget == current_widget:
                tab = t
                break
        
        if not tab or tab.is_readonly:
            return False
        
        # 获取编辑器内容
        if isinstance(current_widget, CodeEditor):
            content = current_widget.toPlainText()
            
            try:
                if self.file_manager:
                    self.file_manager.update_file(tab.path, content)
                else:
                    with open(tab.path, 'w', encoding='utf-8') as f:
                        f.write(content)
                
                # 更新修改状态
                current_widget.set_modified(False)
                tab.is_modified = False
                
                # 更新标签页标题（移除修改标记）
                self._update_tab_title(tab.path)
                
                # 发出保存信号
                self.file_saved.emit(tab.path)
                
                if self.logger:
                    self.logger.info(f"File saved: {tab.path}")
                
                return True
                
            except Exception as e:
                if self.logger:
                    self.logger.error(f"Failed to save file: {tab.path}, error: {e}")
                return False
        
        return False

    def save_all_files(self) -> int:
        """
        保存所有已修改的文件（全局保存）
        
        Returns:
            int: 成功保存的文件数量
        """
        saved_count = 0
        
        for path, tab in self._tabs.items():
            # 跳过只读文件
            if tab.is_readonly:
                continue
            
            # 检查是否有修改
            if isinstance(tab.widget, CodeEditor) and tab.widget.is_modified():
                content = tab.widget.toPlainText()
                
                try:
                    if self.file_manager:
                        self.file_manager.update_file(path, content)
                    else:
                        with open(path, 'w', encoding='utf-8') as f:
                            f.write(content)
                    
                    # 更新修改状态
                    tab.widget.set_modified(False)
                    tab.is_modified = False
                    
                    # 更新标签页标题（移除修改标记）
                    self._update_tab_title(path)
                    
                    # 发出保存信号
                    self.file_saved.emit(path)
                    
                    saved_count += 1
                    
                    if self.logger:
                        self.logger.info(f"File saved: {path}")
                        
                except Exception as e:
                    if self.logger:
                        self.logger.error(f"Failed to save file: {path}, error: {e}")
        
        if self.logger and saved_count > 0:
            self.logger.info(f"Save all completed: {saved_count} file(s) saved")
        
        return saved_count

    def reset_all_modification_states(self):
        """
        重置所有打开文件的修改状态
        
        用于会话恢复后，确保所有文件显示为未修改状态
        """
        for path, tab in self._tabs.items():
            if tab.is_readonly:
                continue
            
            if isinstance(tab.widget, CodeEditor):
                # 重置 QTextDocument 和 CodeEditor 的修改状态
                tab.widget.document().setModified(False)
                tab.widget.set_modified(False)
                tab.is_modified = False
                
                # 更新标签页标题（移除修改标记）
                self._update_tab_title(path)
        
        if self.logger:
            self.logger.debug("All modification states reset")

    def undo(self) -> None:
        """
        编辑器级别撤销
        
        调用当前编辑器的内置撤销功能（QPlainTextEdit.undo）
        """
        current_widget = self._tab_widget.currentWidget()
        if isinstance(current_widget, CodeEditor):
            current_widget.undo()

    def redo(self) -> None:
        """
        编辑器级别重做
        
        调用当前编辑器的内置重做功能（QPlainTextEdit.redo）
        """
        current_widget = self._tab_widget.currentWidget()
        if isinstance(current_widget, CodeEditor):
            current_widget.redo()

    def can_undo(self) -> bool:
        """检查是否可以撤销"""
        current_widget = self._tab_widget.currentWidget()
        if isinstance(current_widget, CodeEditor):
            return current_widget.document().isUndoAvailable()
        return False

    def can_redo(self) -> bool:
        """检查是否可以重做"""
        current_widget = self._tab_widget.currentWidget()
        if isinstance(current_widget, CodeEditor):
            return current_widget.document().isRedoAvailable()
        return False

    def get_content(self) -> Optional[str]:
        """
        获取当前编辑器内容
        
        Returns:
            str: 编辑器内容，无内容返回 None
        """
        current_widget = self._tab_widget.currentWidget()
        if isinstance(current_widget, CodeEditor):
            return current_widget.toPlainText()
        return None

    def set_readonly(self, readonly: bool):
        """
        设置只读模式
        
        Args:
            readonly: 是否只读
        """
        self._is_readonly_mode = readonly
        
        # 更新所有代码编辑器的只读状态
        for tab in self._tabs.values():
            if isinstance(tab.widget, CodeEditor):
                tab.widget.setReadOnly(readonly or tab.is_readonly)
        
        # 更新只读标签显示
        if readonly:
            self._readonly_label.setText(self._get_text("status.readonly", "READ ONLY"))
            self._readonly_label.show()
        else:
            self._readonly_label.hide()

    def open_tab(self, path: str) -> bool:
        """
        在新标签页打开文件
        
        Args:
            path: 文件路径
            
        Returns:
            bool: 是否成功
        """
        return self.load_file(path)

    def close_tab(self, index: int) -> bool:
        """
        关闭指定标签页
        
        Args:
            index: 标签页索引
            
        Returns:
            bool: 是否成功
        """
        if index < 0 or index >= self._tab_widget.count():
            return False
        
        widget = self._tab_widget.widget(index)
        
        # 查找对应的标签页数据
        path_to_remove = None
        for path, tab in self._tabs.items():
            if tab.widget == widget:
                path_to_remove = path
                break
        
        if path_to_remove:
            tab = self._tabs[path_to_remove]
            
            # 检查是否有未保存的修改
            if isinstance(widget, CodeEditor) and widget.is_modified():
                reply = QMessageBox.question(
                    self,
                    self._get_text("dialog.confirm.title", "Confirm"),
                    f"Save changes to {os.path.basename(path_to_remove)}?",
                    QMessageBox.StandardButton.Save |
                    QMessageBox.StandardButton.Discard |
                    QMessageBox.StandardButton.Cancel
                )
                
                if reply == QMessageBox.StandardButton.Save:
                    self.save_file()
                elif reply == QMessageBox.StandardButton.Cancel:
                    return False
            
            # 移除标签页
            self._tab_widget.removeTab(index)
            del self._tabs[path_to_remove]
            
            # 更新空状态
            self._update_empty_state()
            
            # 通知可编辑文件状态变化（用于禁用保存按钮）
            self._emit_editable_file_state()
            
            if self.logger:
                self.logger.debug(f"Tab closed: {path_to_remove}")
            
            return True
        
        return False

    def close_all_tabs(self):
        """关闭所有标签页"""
        while self._tab_widget.count() > 0:
            self.close_tab(0)

    def get_open_files(self) -> list:
        """获取当前打开的所有文件路径列表"""
        return list(self._tabs.keys())

    def get_current_file(self) -> Optional[str]:
        """获取当前激活的文件路径"""
        current_widget = self._tab_widget.currentWidget()
        if not current_widget:
            return None
        
        for path, tab in self._tabs.items():
            if tab.widget == current_widget:
                return path
        return None

    def switch_to_file(self, path: str) -> bool:
        """切换到指定文件的标签页"""
        if path not in self._tabs:
            return False
        
        tab = self._tabs[path]
        index = self._tab_widget.indexOf(tab.widget)
        if index >= 0:
            self._tab_widget.setCurrentIndex(index)
            return True
        return False

    def _close_current_tab(self):
        """关闭当前标签页"""
        current_index = self._tab_widget.currentIndex()
        if current_index >= 0:
            self.close_tab(current_index)

    def _emit_undo_redo_state(self):
        """发出撤销/重做状态变化信号"""
        can_undo = self.can_undo()
        can_redo = self.can_redo()
        self.undo_redo_state_changed.emit(can_undo, can_redo)

    def _emit_editable_file_state(self):
        """发出可编辑文件状态变化信号"""
        has_editable = any(
            not tab.is_readonly for tab in self._tabs.values()
        )
        self.editable_file_state_changed.emit(has_editable)

    def _on_undo_available_changed(self, available: bool):
        """撤销可用状态变化"""
        self._emit_undo_redo_state()

    def _on_redo_available_changed(self, available: bool):
        """重做可用状态变化"""
        self._emit_undo_redo_state()

    def _on_editor_modification_changed(self, path: str, modified: bool):
        """编辑器修改状态变化处理"""
        # 更新标签页标题（显示/隐藏未保存标记）
        self._update_tab_title(path)


    # ============================================================
    # 事件处理
    # ============================================================

    def _on_tab_close_requested(self, index: int):
        """标签页关闭请求"""
        self.close_tab(index)

    def _on_current_tab_changed(self, index: int):
        """当前标签页变化"""
        if index < 0:
            # 无标签页时，禁用撤销/重做
            self.undo_redo_state_changed.emit(False, False)
            return
        
        widget = self._tab_widget.widget(index)
        
        # 查找对应的标签页数据
        for path, tab in self._tabs.items():
            if tab.widget == widget:
                self._update_status_bar(path)
                break
        
        # 更新撤销/重做状态
        self._emit_undo_redo_state()
        
        # 更新滚动按钮状态
        self._update_scroll_buttons()

    def _on_scroll_left(self):
        """向左滚动（切换到前一个标签）"""
        current = self._tab_widget.currentIndex()
        if current > 0:
            self._tab_widget.setCurrentIndex(current - 1)

    def _on_scroll_right(self):
        """向右滚动（切换到后一个标签）"""
        current = self._tab_widget.currentIndex()
        count = self._tab_widget.count()
        if current < count - 1:
            self._tab_widget.setCurrentIndex(current + 1)

    def _update_scroll_buttons(self):
        """更新滚动按钮的启用状态"""
        tab_count = self._tab_widget.count()
        current_index = self._tab_widget.currentIndex()
        
        # 有多个标签时显示按钮
        show_buttons = tab_count > 1
        self._scroll_left_btn.setVisible(show_buttons)
        self._scroll_right_btn.setVisible(show_buttons)
        
        if show_buttons:
            self._scroll_left_btn.setEnabled(current_index > 0)
            self._scroll_right_btn.setEnabled(current_index < tab_count - 1)

    def _on_tab_context_menu(self, position: QPoint):
        """标签页右键菜单"""
        # 获取点击的标签页索引
        tab_bar = self._tab_widget.tabBar()
        index = tab_bar.tabAt(position)
        
        if index < 0:
            return
        
        # 创建菜单
        menu = QMenu(self)
        
        # 关闭
        close_action = QAction(self._get_text("btn.close", "Close"), self)
        close_action.triggered.connect(lambda: self.close_tab(index))
        menu.addAction(close_action)
        
        # 关闭其他
        close_others_action = QAction("Close Others", self)
        close_others_action.triggered.connect(lambda: self._close_other_tabs(index))
        menu.addAction(close_others_action)
        
        # 关闭全部
        close_all_action = QAction("Close All", self)
        close_all_action.triggered.connect(self.close_all_tabs)
        menu.addAction(close_all_action)
        
        menu.addSeparator()
        
        # 复制路径
        copy_path_action = QAction(self._get_text("file_browser.copy_path", "Copy Path"), self)
        copy_path_action.triggered.connect(lambda: self._copy_tab_path(index))
        menu.addAction(copy_path_action)
        
        # 显示菜单
        menu.exec(tab_bar.mapToGlobal(position))

    def _close_other_tabs(self, keep_index: int):
        """关闭其他标签页"""
        # 从后往前关闭，避免索引变化问题
        for i in range(self._tab_widget.count() - 1, -1, -1):
            if i != keep_index:
                self.close_tab(i)

    def _copy_tab_path(self, index: int):
        """复制标签页文件路径"""
        widget = self._tab_widget.widget(index)
        for path, tab in self._tabs.items():
            if tab.widget == widget:
                clipboard = QApplication.clipboard()
                clipboard.setText(path)
                if self.logger:
                    self.logger.debug(f"Path copied: {path}")
                break

    def _update_cursor_position(self):
        """更新光标位置显示"""
        current_widget = self._tab_widget.currentWidget()
        if isinstance(current_widget, CodeEditor):
            line, col = current_widget.get_cursor_position()
            self._line_col_label.setText(f"Ln {line}, Col {col}")

    def _update_tab_title(self, path: str):
        """更新标签页标题"""
        if path not in self._tabs:
            return
        
        tab = self._tabs[path]
        index = self._tab_widget.indexOf(tab.widget)
        
        if index < 0:
            return
        
        file_name = os.path.basename(path)
        
        # 检查是否有未保存的修改
        if isinstance(tab.widget, CodeEditor) and tab.widget.is_modified():
            self._tab_widget.setTabText(index, f"{file_name} ●")
        else:
            self._tab_widget.setTabText(index, file_name)

    def _update_status_bar(self, path: str):
        """更新状态栏"""
        ext = os.path.splitext(path)[1].lower()
        
        # 更新文件类型
        file_type_map = {
            '.cir': 'SPICE',
            '.sp': 'SPICE',
            '.spice': 'SPICE',
            '.json': 'JSON',
            '.txt': 'Plain Text',
            '.md': 'Markdown',
            '.markdown': 'Markdown',
            '.docx': 'Word Document',
            '.pdf': 'PDF Document',
            '.png': 'PNG Image',
            '.jpg': 'JPEG Image',
            '.jpeg': 'JPEG Image',
        }
        
        file_type = file_type_map.get(ext, 'Plain Text')
        self._file_type_label.setText(file_type)
        
        # 更新光标位置
        self._update_cursor_position()

    # ============================================================
    # 国际化支持
    # ============================================================

    def retranslate_ui(self):
        """刷新所有 UI 文本（语言切换时调用）"""
        # 空状态提示
        for child in self._empty_widget.findChildren(QLabel):
            if child.property("empty_hint"):
                child.setText(self._get_text("hint.select_file", "Select a file to view"))
        
        # 打开工作区按钮
        if hasattr(self, '_open_workspace_btn') and self._open_workspace_btn:
            self._open_workspace_btn.setText(self._get_text("btn.open_workspace", "Open Workspace"))
        
        # 只读标签
        if self._readonly_label.isVisible():
            self._readonly_label.setText(self._get_text("status.readonly", "READ ONLY"))

    # ============================================================
    # 事件订阅
    # ============================================================

    def _subscribe_events(self):
        """订阅事件"""
        if self.event_bus:
            from shared.event_types import (
                EVENT_LANGUAGE_CHANGED,
                EVENT_STATE_PROJECT_OPENED,
                EVENT_STATE_PROJECT_CLOSED,
                EVENT_WORKFLOW_LOCKED,
                EVENT_WORKFLOW_UNLOCKED
            )
            
            self.event_bus.subscribe(EVENT_LANGUAGE_CHANGED, self._on_language_changed)
            self.event_bus.subscribe(EVENT_STATE_PROJECT_OPENED, self._on_project_opened)
            self.event_bus.subscribe(EVENT_STATE_PROJECT_CLOSED, self._on_project_closed)
            self.event_bus.subscribe(EVENT_WORKFLOW_LOCKED, self._on_workflow_locked)
            self.event_bus.subscribe(EVENT_WORKFLOW_UNLOCKED, self._on_workflow_unlocked)

    def _on_language_changed(self, event_data: Dict[str, Any]):
        """语言变更事件处理"""
        self.retranslate_ui()
        if self.logger:
            new_lang = event_data.get("new_language", "unknown")
            self.logger.debug(f"Code editor language changed to: {new_lang}")

    def _on_project_opened(self, event_data: Dict[str, Any]):
        """项目打开事件处理"""
        # 关闭所有标签页，准备加载新项目文件
        self.close_all_tabs()
        # 更新空状态显示（隐藏打开工作区按钮）
        self._update_empty_state()
        if self.logger:
            # 业务数据在 event_data["data"] 中
            data = event_data.get("data", {})
            project_path = data.get("path", "unknown") if isinstance(data, dict) else "unknown"
            self.logger.info(f"Code editor cleared for new project: {project_path}")

    def _on_project_closed(self, event_data: Dict[str, Any]):
        """项目关闭事件处理"""
        # 关闭所有标签页，清空编辑器内容
        self.close_all_tabs()
        # 更新空状态显示（显示打开工作区按钮）
        self._update_empty_state()
        if self.logger:
            self.logger.info("Code editor cleared due to project close")

    def _on_workflow_locked(self, event_data: Dict[str, Any]):
        """工作流锁定事件处理"""
        self.set_readonly(True)
        if self.logger:
            self.logger.debug("Code editor set to readonly mode (workflow locked)")

    def _on_workflow_unlocked(self, event_data: Dict[str, Any]):
        """工作流解锁事件处理"""
        self.set_readonly(False)
        if self.logger:
            self.logger.debug("Code editor readonly mode disabled (workflow unlocked)")


# ============================================================
# 模块导出
# ============================================================

__all__ = [
    "CodeEditorPanel",
    "CodeEditor",
    "SpiceHighlighter",
    "JsonHighlighter",
    "ImageViewer",
    "DocumentViewer",
    "EDITABLE_EXTENSIONS",
    "IMAGE_EXTENSIONS",
    "DOCUMENT_EXTENSIONS",
]
